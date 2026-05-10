#!/usr/bin/env python3
"""Convert Markdown Hardware Guide to DOCX with professional formatting."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Create document
doc = Document()

# Set up styles
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

# Colors
HEADER_COLOR = RGBColor(46, 117, 182)  # #2E75B6
ACCENT_COLOR = RGBColor(46, 117, 182)

def add_heading(doc, text, level=1):
    """Add formatted heading."""
    heading = doc.add_heading(text, level=level)
    if level == 1:
        heading.style = 'Heading 1'
        for run in heading.runs:
            run.font.color.rgb = HEADER_COLOR
            run.font.bold = True
            run.font.size = Pt(20)
    elif level == 2:
        heading.style = 'Heading 2'
        for run in heading.runs:
            run.font.color.rgb = ACCENT_COLOR
            run.font.size = Pt(14)
    return heading

def add_bullet_point(doc, text, level=0):
    """Add bullet point."""
    p = doc.add_paragraph(text, style=f'List Bullet {level + 1}')
    for run in p.runs:
        run.font.size = Pt(11)

def add_numbered_item(doc, text, level=0):
    """Add numbered item."""
    p = doc.add_paragraph(text, style=f'List Number {level + 1}')
    for run in p.runs:
        run.font.size = Pt(11)

def shade_cell(cell, color):
    """Shade table cell."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

# Title Page
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run('STALLWACHE')
title_run.font.size = Pt(36)
title_run.font.bold = True
title_run.font.color.rgb = HEADER_COLOR

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.add_run('Hardware-Setup-Guide')
subtitle_run.font.size = Pt(28)
subtitle_run.font.color.rgb = HEADER_COLOR

subsubtitle = doc.add_paragraph()
subsubtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subsubtitle_run = subsubtitle.add_run('Rollei Safetycam HD 20 Integration')
subsubtitle_run.font.size = Pt(18)
subsubtitle_run.font.bold = True

desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
desc_run = desc.add_run('KI-basiertes Kalberkennungs- und Stallüberwachungssystem')
desc_run.font.size = Pt(12)
desc_run.font.italic = True

footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer.add_run('Version 1.0 | Mai 2026')
footer_run.font.size = Pt(10)
footer_run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_page_break()

# 1. Systemübersicht
add_heading(doc, '1. Systemübersicht', 1)
doc.add_paragraph('Das Stallwache-System besteht aus drei Hauptkomponenten:')
add_bullet_point(doc, 'Rollei Safetycam HD 20 – Netzwerk-Überwachungskamera')
add_bullet_point(doc, 'Edge-Server (lokaler PC/Laptop) – KI-Inferenz und Alert-Verarbeitung')
add_bullet_point(doc, 'Rollei Cloud & Messenger – Backup-Speicher und Benachrichtigungen')

# 2. Technische Spezifikationen
add_heading(doc, '2. Rollei Safetycam HD 20 – Technische Spezifikationen', 1)

add_heading(doc, '2.1 Hardware-Specs', 2)
table = doc.add_table(rows=11, cols=2)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Parameter'
hdr_cells[1].text = 'Wert'
for cell in hdr_cells:
    shade_cell(cell, '2E75B6')
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

specs = [
    ['Auflösung', '1920×1080 Pixel (Full HD)'],
    ['Framerate', '30 fps bei Full HD'],
    ['Sensor', '1/3″ CMOS'],
    ['Blickwinkel', '160° diagonal'],
    ['Nachtsicht', 'Infrarot-LEDs (bis 15m)'],
    ['Stromversorgung', '12V DC / PoE'],
    ['Wasserschutz', 'IP66'],
    ['Betriebstemp.', '-10°C bis +50°C'],
    ['Gewicht', 'ca. 280g'],
    ['Abmessungen', '180×108×86 mm']
]

for i, (param, value) in enumerate(specs, start=1):
    row_cells = table.rows[i].cells
    row_cells[0].text = param
    row_cells[1].text = value

doc.add_paragraph()

add_heading(doc, '2.2 Netzwerk-Features', 2)
add_bullet_point(doc, 'RTSP-Stream: rtsp://[CAMERA_IP]:554/stream')
add_bullet_point(doc, 'ONVIF-Support: Standardisierte Geräte-Integration')
add_bullet_point(doc, 'HTTP-API: Web-Interface für Konfiguration')
add_bullet_point(doc, 'WiFi 802.11ac: 5GHz/2.4GHz oder Gigabit Ethernet')

doc.add_page_break()

# 3. Netzwerk-Konfiguration
add_heading(doc, '3. Netzwerk-Konfiguration', 1)

add_heading(doc, '3.1 Kamera-Initialisierung', 2)
add_numbered_item(doc, 'Stromversorgung anschließen (12V DC oder PoE-Switch)')
add_numbered_item(doc, 'Warten Sie 30 Sekunden auf Boot')
add_numbered_item(doc, 'LED sollte blau/grün blinken')
add_numbered_item(doc, 'Kamera versucht sich über DHCP einzubuchen')

add_heading(doc, '3.2 IP-Adresse ermitteln', 2)
add_heading(doc, 'Methode A: Rollei Setup-App', 3)
add_numbered_item(doc, 'Laden Sie die "Rollei SafetyCam" App herunter')
add_numbered_item(doc, 'App scannt lokales Netzwerk nach Kamera')
add_numbered_item(doc, 'Notieren Sie die IP-Adresse (z.B. 192.168.1.100)')

add_heading(doc, 'Methode B: Router-Admin-Panel', 3)
add_numbered_item(doc, 'Öffnen Sie Router-Admin-Panel (192.168.1.1 oder 192.168.0.1)')
add_numbered_item(doc, 'Navigieren Sie zu "Verbundene Geräte"')
add_numbered_item(doc, 'Suchen Sie nach "Rollei" oder MAC 00:1A:AB')
add_numbered_item(doc, 'Notieren Sie die IP-Adresse')

add_heading(doc, '3.3 Statische IP zuweisen', 2)
add_numbered_item(doc, 'Öffnen Sie Web-Interface: http://[CAMERA_IP]')
add_numbered_item(doc, 'Standard-Login: admin / 12345')
add_numbered_item(doc, 'Gehen Sie zu Einstellungen → Netzwerk')
add_numbered_item(doc, 'Wählen Sie "Statische IP" statt DHCP')
add_numbered_item(doc, 'Empfehlung: 192.168.1.200')

add_heading(doc, '3.4 RTSP-Stream testen', 2)
doc.add_paragraph('Öffnen Sie VLC Media Player:')
add_numbered_item(doc, 'Datei → Stream öffnen')
add_numbered_item(doc, 'URL: rtsp://admin:12345@192.168.1.200:554/stream')
add_numbered_item(doc, 'Video startet ✓ – RTSP funktioniert')

doc.add_page_break()

# 4. Platzierung
add_heading(doc, '4. Kamera-Platzierung im Stall', 1)

add_heading(doc, '4.1 Optimale Positionierung', 2)
add_bullet_point(doc, 'Höhe: 2–3m')
add_bullet_point(doc, 'Winkel: 45° nach unten')
add_bullet_point(doc, 'Entfernung: 2–5m zum Abkalbe-Bereich')
add_bullet_point(doc, 'Vermeiden Sie: Gegenlicht, Reflektionen')

add_heading(doc, '4.2 Stromversorgung', 2)
add_bullet_point(doc, 'Separates 12V-Netzteil oder PoE-Injektor')
add_bullet_point(doc, 'Kabelführung: entlang Stallwand')
add_bullet_point(doc, 'Netzwerk-Kabel: Cat5e/Cat6, min. 10m')

# 5. Edge-Server
add_heading(doc, '5. Edge-Server Anforderungen', 1)

add_heading(doc, '5.1 Mindestanforderungen', 2)
table2 = doc.add_table(rows=9, cols=2)
table2.style = 'Light Grid Accent 1'
hdr_cells2 = table2.rows[0].cells
hdr_cells2[0].text = 'Komponente'
hdr_cells2[1].text = 'Anforderung'
for cell in hdr_cells2:
    shade_cell(cell, '2E75B6')
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

specs2 = [
    ['CPU', 'Intel i5/i7 oder AMD Ryzen 5/7 (4+ Cores)'],
    ['RAM', '8 GB (16 GB empfohlen)'],
    ['GPU', 'Optional: NVIDIA RTX 3060+'],
    ['Speicher', '250 GB SSD'],
    ['OS', 'Windows 10/11, Linux, macOS'],
    ['Netzwerk', 'Gigabit Ethernet oder 5GHz WiFi'],
    ['RAM-Nutzung', '~2–4 GB während Inferenz'],
    ['CPU-Last', '~30–60% bei 30fps 1080p']
]

for i, (comp, req) in enumerate(specs2, start=1):
    row_cells = table2.rows[i].cells
    row_cells[0].text = comp
    row_cells[1].text = req

add_heading(doc, '5.2 Software-Setup', 2)
add_numbered_item(doc, 'Python 3.10+ installieren')
add_numbered_item(doc, 'Dependencies: opencv-python, ultralytics, python-telegram-bot')

doc.add_page_break()

# 7. Sicherheit
add_heading(doc, '7. Sicherheit & Datenschutz', 1)

add_heading(doc, '7.1 Kamera-Sicherheit', 2)
add_bullet_point(doc, 'Standard-Passwort ändern: admin → sicheres Passwort')
add_bullet_point(doc, 'Firewall-Regel: Kamera nur im lokalen Netzwerk')
add_bullet_point(doc, 'Firmware regelmäßig aktualisieren')

add_heading(doc, '7.2 Datenschutz', 2)
add_bullet_point(doc, 'Lokal: 7 Tage, Cloud: 30 Tage')
add_bullet_point(doc, 'Nur Tier-Bilder werden analysiert')

# 8. Troubleshooting
add_heading(doc, '8. Troubleshooting', 1)
table3 = doc.add_table(rows=5, cols=2)
table3.style = 'Light Grid Accent 1'
hdr_cells3 = table3.rows[0].cells
hdr_cells3[0].text = 'Problem'
hdr_cells3[1].text = 'Lösung'
for cell in hdr_cells3:
    shade_cell(cell, '2E75B6')
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

ts = [
    ['Kamera antwortet nicht', 'Stromversorgung prüfen, Router neu starten'],
    ['RTSP-Stream fehlt', 'Netzwerk-Stabilität überprüfen, WiFi → Ethernet'],
    ['Niedriges FPS', 'CPU/RAM überprüfen, Qualität reduzieren'],
    ['Cloud-Upload fehlt', 'Cloud-Login prüfen, Internet-Verbindung testen']
]

for i, (prob, sol) in enumerate(ts, start=1):
    row_cells = table3.rows[i].cells
    row_cells[0].text = prob
    row_cells[1].text = sol

doc.add_page_break()

# Final Section
add_heading(doc, 'Nächste Schritte', 1)
add_numbered_item(doc, 'Konsultieren Sie den Python-Code-Guide')
add_numbered_item(doc, 'Docker-Container aufsetzen')
add_numbered_item(doc, 'KI-Modelle mit Testvideos prüfen')
add_numbered_item(doc, 'Alert-Schwellenwerte tunen')

doc.add_paragraph()
footer_p = doc.add_paragraph('Support: stallwache123@gmail.com')
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in footer_p.runs:
    run.font.italic = True
    run.font.size = Pt(10)

# Save
doc.save('C:\\Users\\axe2k\\Desktop\\Projekt Stall\\stallwache\\Stallwache\\Rollei_HD20_Hardware_Setup_Guide.docx')
print("✓ DOCX document created successfully!")
